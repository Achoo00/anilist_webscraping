# Done
# Scrapes anilist for titles, start and finish dates
# Outputs data in a doc formatted for anilist challenges

from bs4 import BeautifulSoup
import requests
import webbrowser
import pyautogui
import time
import win32clipboard
import win32api
import win32con
import sys
from docx import Document
import datetime
import os
import pytesseract
import PIL.Image

url = ''
urls = []
title = ''
titles = []
dates = []
start_date = ''
start_dates = []
finish_date = ''
finish_dates = []
site_url = ''
site_urls = []
pyautogui.FAILSAFE = True


def screenshot():
    shot = pyautogui.screenshot(region=(500, 620, 500, 60))
    now = datetime.datetime.now()
    timenow = now.strftime("%H_%M_%S")
    path = "C://Users//amaha//PycharmProjects//anilist_webscraping//img//" + str(datetime.date.today())
    os.chdir("C:\\Users\\amaha\\PycharmProjects\\anilist_webscraping\\img\\" + str(datetime.date.today()))
    try:
        shot.save(path + '//' + timenow + '.png')
    except FileNotFoundError:
        os.makedirs(path)
        shot.save(path + '//' + timenow + '.png')

    pytesseract.pytesseract.tesseract_cmd = r"C:\Users\amaha\AppData\Local\Tesseract-OCR\tesseract.exe"

    myconfig = r"--psm 3 --oem 3"
    text = pytesseract.image_to_string(PIL.Image.open(f"{timenow}.png"), config=myconfig)
    new_text = text.split()
    new_text.sort()
    for dash in new_text:
        if '-' in dash:
            dates.append(dash)


def click():
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTDOWN, 0, 0)
    time.sleep(0.1)
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTUP, 0, 0)
    time.sleep(0.1)


def alt_tab():
    pyautogui.keyDown('alt')
    time.sleep(.2)
    pyautogui.press('tab')
    time.sleep(.2)
    pyautogui.keyUp('alt')


def output_text():
    n = len(titles)
    document = Document()
    start_dates = dates[::2]
    finish_dates = dates[1::2]
    try:
        for i in range(n):
            print(start_dates[i])
            document.add_paragraph(f"[{titles[i].strip()}]({urls[i].strip()})")
            document.add_paragraph(f"Start: {start_dates[i].strip()} End: {finish_dates[i].strip()}")
            i += 1
    except IndexError:
        for z in range(n):
            document.add_paragraph(f"[{titles[z].strip()}]({urls[z].strip()})")
            z += 1
    document.save('output.docx')


def anime():
    length = len(site_urls)
    open(length)

def open(x):
    for y in range(x):
        webbrowser.open_new(site_urls[y])
        os.chdir("/anilist_webscraping")
        locate(site_urls[y])
        y += 1
    output_text()
    os.system('start output.docx')
    sys.exit()

def locate(url):
    while True:
        completed = pyautogui.locateCenterOnScreen("completed_button.png", region=(285, 450, 300, 600),
                                                   grayscale=False, confidence=0.7)
        #region=(0, 0, 1920, 1080)
        #region=(290, 660, 180, 65)
        print("Searching")
        if completed is not None:
            time.sleep(1)
            pyautogui.moveTo(completed)
            time.sleep(0.1)
            click()
            time.sleep(1)
            screenshot()
            time.sleep(0.1)
            html_text = requests.get(url).text
            soup = BeautifulSoup(html_text, 'lxml')
            title = soup.find('h1', class_='').text
            titles.append(title)
            #print("title")
            break
            #alt_tab()

def character_voice_actor():
    length = len(site_urls)
    character_voice_actor_title(length)
    character_voice_actor_title(length)

def character_voice_actor_title(i):
    for z in range(i):
        html_text = requests.get(site_urls[z]).text
        soup = BeautifulSoup(html_text, 'lxml')
        title = soup.find('h1', class_="name").text
        titles.append(title)
    output_text()
    print(dates)
    os.system('start output.docx')
    sys.exit()


def anime_append():
    site_url = input("Site URL:")
    if 'https://anilist.co/' in site_url:
        site_urls.append(site_url)
        urls.append(site_url)
        anime_append()
    elif site_url == 'done':
        anime()
    else:
        retry()

def character_voice_actor_append():
    site_url = input("Site URL:")
    if 'https://anilist.co/' in site_url:
        site_urls.append(site_url)
        urls.append(site_url)
        character_voice_actor_append()
    elif site_url == 'done':
        character_voice_actor()
    else:
        retry()

def retry():
    print("Please try again")
    selecting()


def selecting():
    select = input("Anime(an), Character(ch) or Voice Actor(vc)?:")
    if select == 'done':
        output_text()
        os.system('start output.docx')
        sys.exit()
    else:
        if select == 'an':
            anime_append()
        elif select == 'ch':
            character_voice_actor_append()
        elif select == 'vc':
            character_voice_actor_append()
        else:
            retry()


if __name__ == '__main__':
    path = "C://Users//amaha//PycharmProjects//anilist_webscraping//img//" + str(datetime.date.today())
    try:
        os.makedirs(path)
    except FileExistsError:
        pass
    selecting()
    sys.exit()

# Samples
# https://anilist.co/anime/1689/5-Centimeters-per-Second/
# https://anilist.co/anime/106286/Weathering-With-You/
