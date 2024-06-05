from docxtpl import DocxTemplate
import docx
from bs4 import BeautifulSoup
import requests
import webbrowser
import pyautogui
import time
import win32clipboard
import win32api
import win32con
import sys
from docx import Document
import subprocess
import os


def click():
    time.sleep(0.1)
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTDOWN, 0, 0)
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTUP, 0, 0)
    time.sleep(0.1)

def double_click():
    time.sleep(0.1)
    pyautogui.doubleClick()
    time.sleep(0.1)


def template():
    doc_template = DocxTemplate("inviteTmpl - Copy.docx")

    context = {
        "O": "X",
        "Anime_Title": "Byousoku 5 Centimeter",
        "Link": "https://anilist.co/anime/1689/5-Centimeters-per-Second/",
        "Start": "Start: 2018-07-17",
        "Finish": "Finish: 2018-07-17"
    }

    doc_template.render(context)

    doc_template.save("challenges/challenge.docx")

def doc():
    doc = docx.Document("inviteTmpl - Copy.docx")
    #print(doc.paragraphs[1].text)
    print(doc.paragraphs[1].runs[1].text)
    test = doc.paragraphs[1].runs[1]
    test = doc.add_paragraph('gura')
    doc.save('inviteTmpl - Copy.docx')

def click_replace():
    print('yo')

def replace():
    pyautogui.hotkey('ctrl', 'h')
    pyautogui.write('[O]')
    tab()
    pyautogui.write('[X]')
    find_replace_all()
    reset()
    pyautogui.write('[Anime_Title](https://anilist.co/anime/00000/)')
    tab()
    pyautogui.write('[{{Anime_Title}}]({{Link}})')
    find_replace_all()
    reset()
    pyautogui.write('Start: YYYY-MM-DD Finish: YYYY-MM-DD')
    tab()
    pyautogui.write('{{Start}} {{Finish}}')
    find_replace_all()
    close()
    pyautogui.hotkey('ctrl', 's')
    time.sleep(1)
    pyautogui.hotkey('enter')
    exit()
    template()

def tab():
    time.sleep(0.1)
    #image = pyautogui.locateCenterOnScreen("filter.png", region=(0, 0, 1920, 1080),
    #                                       grayscale=False, confidence=0.9)
    #if image is not None:
    pyautogui.moveTo(1233,448)
    time.sleep(0.1)
    click()

def find_replace_all():
    time.sleep(0.1)
    image = pyautogui.locateCenterOnScreen("replace_all.png", region=(0, 0, 1920, 1080),
                                                       grayscale=False, confidence=0.9)
    if image is not None:
        pyautogui.moveTo(image)
        time.sleep(0.1)
        click()

def reset():
    time.sleep(0.1)
    #image = pyautogui.locateCenterOnScreen("reset.png", region=(0, 0, 1920, 1080),
    #                                                   grayscale=False, confidence=0.8)
    #if image is not None:
    pyautogui.moveTo(1233,374)
    time.sleep(0.1)
    click()

def close():
    time.sleep(0.1)
    #image = pyautogui.locateCenterOnScreen("reset.png", region=(0, 0, 1920, 1080),
    #                                                   grayscale=False, confidence=0.8)
    #if image is not None:
    pyautogui.moveTo(1460,726)
    time.sleep(0.1)
    click()

def exit():
    time.sleep(0.1)
    #image = pyautogui.locateCenterOnScreen("reset.png", region=(0, 0, 1920, 1080),
    #                                                   grayscale=False, confidence=0.8)
    #if image is not None:
    pyautogui.moveTo(1892,3)
    time.sleep(0.1)
    click()


time.sleep(3)
replace()