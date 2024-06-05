# Scrapes anilist for titles, start and finish dates
# Outputs data in a doc formatted for anilist challenges

from bs4 import BeautifulSoup
import requests
import webbrowser
import pyautogui
import time
import win32clipboard
import win32api
import win32con
import sys
from docx import Document
import os

url = ''
urls = []
title = ''
titles = []
start_date = ''
start_dates = []
finish_date = ''
finish_dates = []

pyautogui.FAILSAFE = True


def triple_click(x, y):
    pyautogui.tripleClick(x, y)
    time.sleep(0.1)
    pyautogui.hotkey('ctrl', 'c')
    win32clipboard.OpenClipboard()
    data = win32clipboard.GetClipboardData()
    win32clipboard.CloseClipboard()
    if x == 591:
        start_date = data
        start_dates.append(start_date)
    else:
        finish_date = data
        finish_dates.append(finish_date)


def click():
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTDOWN, 0, 0)
    win32api.mouse_event(win32con.MOUSEEVENTF_LEFTUP, 0, 0)
    time.sleep(0.1)


def alt_tab():
    pyautogui.keyDown('alt')
    time.sleep(.2)
    pyautogui.press('tab')
    time.sleep(.2)
    pyautogui.keyUp('alt')
    selecting()


def output_text():
    n = len(titles)
    document = Document()
    for i in range(n):
        document.add_paragraph(f"[{titles[i].strip()}]({urls[i].strip()})")
        i += 1
    document.save('word.docx')


def output_text_2():
    n = len(titles)
    document = Document()
    for i in range(n):
        document.add_paragraph(f"[{titles[i].strip()}]({urls[i].strip()})")
        document.add_paragraph(f"Start: {start_dates[i].strip()} End: {finish_dates[i].strip()}")
        i += 1
    document.save('word.docx')


def anime(url):
    urls.append(url)
    webbrowser.open_new(url)
    while True:
        completed = pyautogui.locateCenterOnScreen("completed_button.png", region=(0, 0, 1920, 1080),
                                                   grayscale=False, confidence=0.75)
        print("Searching")
        if completed is not None:
            time.sleep(2)
            pyautogui.moveTo(completed)
            time.sleep(0.1)
            click()
            triple_click(591, 647)
            time.sleep(0.1)
            triple_click(847, 646)
            html_text = requests.get(url).text
            soup = BeautifulSoup(html_text, 'lxml')
            title = soup.find('h1', class_='').text
            titles.append(title)
            alt_tab()


def character_voice_actor(url):
    html_text = requests.get(url).text
    urls.append(url)
    soup = BeautifulSoup(html_text, 'lxml')
    title = soup.find('h1', class_="name").text
    titles.append(title)
    selecting()


def selecting():
    select = input("Anime(an), Character(ch) or Voice Actor(vc)?:")
    if select == 'done':
        output_text_2()
        os.system('start word.docx')
        sys.exit()
    else:
        if select == 'an':
            site_url = input("Site URL:")
            anime(site_url)
        elif select == 'ch':
            site_url = input("Site URL:")
            character_voice_actor(site_url)
        elif select == 'vc':
            site_url = input("Site URL:")
            character_voice_actor(site_url)


selecting()
sys.exit()
# Samples
# https://anilist.co/anime/1689/5-Centimeters-per-Second/
# https://anilist.co/anime/106286/Weathering-With-You/
